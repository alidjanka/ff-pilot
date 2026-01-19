import streamlit as st
from docx import Document
from io import BytesIO

from RAG.openai_rag import OpenAIRAG

import datetime
import os
import asyncio
import time

from prompts.prompts import Sections
from utils.document import fill_flb_document
from utils.file_system import set_configuration_files

MAX_PROJEKTE=30

async def generate_document(sections):
    rag = OpenAIRAG(projekt_bezeichnung=st.session_state["projektbezeichnung"])
    full_doc = await rag.build_document_text(sections)
    return full_doc

def ask(query, projekt_bezeichnung):
    rag = OpenAIRAG(projekt_bezeichnung)
    answer, file_names = rag.query(query)
    return answer, file_names

def ask_stream(query: str, projekt_bezeichnung: str):
    rag = OpenAIRAG(projekt_bezeichnung)

    for chunk, files in rag.query_stream(query):
        yield chunk, files

def md_to_docx(md_text: str) -> bytes:
    """
    Convert markdown or plain text into a .docx file (simple formatting).
    Returns raw bytes, ready for Streamlit download.
    """
    doc = Document()
    lines = md_text.split("\n")

    for line in lines:
        if line.startswith("# "):        # H1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):     # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):    # H3
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

st.set_page_config(page_title="Dokument Generator", layout="wide")

# ------------------- Initialize session state -------------------
if "generated_doc" not in st.session_state:
    st.session_state.generated_doc = None

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ------------------- Layout -------------------
st.image("assets/logo.png", width=800)
#st.title("Mogli - FF Pilot")

# Two columns for split screen
left_col, right_col = st.columns([1, 1])

# ------------------- LEFT COLUMN: DOCUMENT GENERATOR -------------------
with left_col:
    st.header("📄 FLB Generator")
    st.markdown(
        "[📄 Vorlage & Masterliste hier](https://drive.google.com/drive/folders/1rrX8hLwrIwzfzdOsyAF4O1TaXfSmhCMc?usp=sharing)"
    )
    if st.button("Aktualisiere Vorlage & Masterliste"):
        with st.spinner("Projektinformationen werden aktualisiert …"):
            try:
                os.remove(st.session_state["masterliste_path"])
                os.remove(st.session_state["template_path"])
            except:
                pass
            time.sleep(3)
            r = set_configuration_files()
            if r is None:
                st.error("Aktualisierung fehlgeschlagen")
            else:
                st.success("Aktualisiert!")
    if len(st.session_state["projects"]) > MAX_PROJEKTE:
        st.warning(f"Maximal erluabte Anzahl von Projekten ist {MAX_PROJEKTE}. Lösche ein altes Projekt und versuch erneut.")
        doc_generation_disabled= True
    elif st.session_state.get('projektbezeichnung'):
        st.subheader(f"Ausgewähltes Projekt: {st.session_state['projektbezeichnung']}")
        doc_generation_disabled = False
    else:
        st.write(f"Noch kein Projekt ausgewählt!")
        doc_generation_disabled= True

    if st.button("📄 FLB generieren", disabled=doc_generation_disabled):     
        #st.session_state.generated_doc = asyncio.run(generate_document(Sections))
        with st.spinner("FLB wird generiert …"):
            try:
                st.session_state.generated_doc = asyncio.run(fill_flb_document(template_path=st.session_state["template_path"], user_inputs={
                    "Projekt":  st.session_state["projektbezeichnung"],
                    "Objektadresse": st.session_state["objektadresse"],
                    "Ansprechpartner": st.session_state["ansprechpartner"]
                }))
                st.success("Dokument bereit zum Download!")
            except:
                st.warning("Es kann einige Zeit dauern, bis Änderungen in der Vorlage und der Masterliste wirksam werden. Bitte warten Sie einen Moment und aktualisieren Sie die Vorlage & Masterliste erneut.")

    st.markdown("---")

    if st.session_state.generated_doc:
        st.subheader("📘 Generiertes Dokument")
        #st.markdown(st.session_state.generated_doc)
        #docx_bytes = md_to_docx(st.session_state.generated_doc)
        st.download_button(
        label="📥 Download",
        data=st.session_state.generated_doc,
        file_name=f"FLB_Repowering_{datetime.date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Noch kein Dokument generiert.")

# ------------------- RIGHT COLUMN: CHAT INTERFACE -------------------
with right_col:
    st.header("💬 Chat mit deinen Unterlagen")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # placeholder at the bottom to scroll to
    scroll_placeholder = st.empty()

    # render chat history
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    streaming = st.session_state.chat_history and st.session_state.chat_history[-1].get("streaming", False)

    # show input only if not streaming
    if not streaming:
        user_input = st.chat_input("Nachricht eingeben...")

        if user_input:
            st.session_state.chat_history.append({
                "role": "user",
                "content": user_input
            })
            st.session_state.chat_history.append({
                "role": "assistant",
                "content": "💡 Schreibt…",
                "streaming": True
            })
            st.rerun()

    # streaming assistant response
    elif streaming:
        last_msg = st.session_state.chat_history[-1]

        with st.chat_message("assistant"):
            placeholder = st.empty()
            full_answer = ""
            cited_files = []

            # start with current content
            placeholder.markdown(last_msg["content"] + "▌")

            for chunk, files in ask_stream(
                st.session_state.chat_history[-2]["content"],
                st.session_state["projektbezeichnung"]
            ):
                if chunk:
                    full_answer += chunk
                    placeholder.markdown(full_answer + "▌")
                    time.sleep(0.01)

                    # auto-scroll by moving placeholder to bottom
                    scroll_placeholder.empty()  # triggers scroll

                if files is not None:
                    cited_files = files

            # final message
            placeholder.markdown(
                f"""{full_answer}

**Gemäß den folgenden Dateien:** {', '.join(cited_files)}
"""
            )

        # update chat history
        last_msg.update({
            "content": full_answer,
            "streaming": False
        })

        st.rerun()



