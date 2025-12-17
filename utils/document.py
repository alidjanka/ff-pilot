from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from io import BytesIO
from RAG.openai_rag import OpenAIRAG
import re

def normalize_instructions(raw: str) -> str:
    text = raw.strip()

    # Replace fancy bullets with simple hyphens
    text = re.sub(r"[•●▪∙]+", "-", text)

    # Remove multiple empty lines
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)

    # Ensure single-line separators become periods
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    cleaned = "\n".join(lines)
    return cleaned

# -------------------------------------------------------
# Helper: insert paragraph AFTER a given paragraph
# -------------------------------------------------------
def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def is_page_break(para):
    for run in para.runs:
        for br in run._element.findall(".//w:br", run._element.nsmap):
            if br.get("{urn:schemas-microsoft-com:office:word}type") == "page":
                return True
    return False

def insert_text_block_after(last_para, text):
    for line in text.split("\n"):
        if line.strip():
            last_para = insert_paragraph_after(last_para, line.strip())

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

# -------------------------------------------------------
# Helper: copy paragraph-level formatting and set cell text
# -------------------------------------------------------
def set_cell_text_preserve_style(src_cell, dest_cell, text):
    """
    Copy paragraph style and formatting from src_cell (usually the label cell)
    to dest_cell, then set dest_cell's text to `text` without destroying layout.
    """
    # Ensure both cells have at least one paragraph
    src_para = src_cell.paragraphs[0] if src_cell.paragraphs else src_cell.add_paragraph()
    dest_para = dest_cell.paragraphs[0] if dest_cell.paragraphs else dest_cell.add_paragraph()

    # Copy paragraph style (if any)
    try:
        dest_para.style = src_para.style
    except Exception:
        # ignore if style can't be copied
        pass

    # Copy paragraph-level alignment and spacing / indents
    try:
        dest_para.alignment = src_para.alignment
    except Exception:
        pass

    try:
        dest_para.paragraph_format.left_indent = src_para.paragraph_format.left_indent
        dest_para.paragraph_format.right_indent = src_para.paragraph_format.right_indent
        dest_para.paragraph_format.first_line_indent = src_para.paragraph_format.first_line_indent
        dest_para.paragraph_format.space_before = src_para.paragraph_format.space_before
        dest_para.paragraph_format.space_after = src_para.paragraph_format.space_after
        dest_para.paragraph_format.line_spacing = src_para.paragraph_format.line_spacing
    except Exception:
        pass

    # Copy cell vertical alignment
    try:
        dest_cell.vertical_alignment = src_cell.vertical_alignment
    except Exception:
        # fallback: do nothing
        pass

    # Clear existing runs in dest_para (without touching paragraph object)
    for run in list(dest_para.runs):
        run.text = ""

    # Add the new text as a single run
    new_run = dest_para.add_run(text)

    # Optionally, copy font size from src first run if available (best-effort)
    try:
        if src_para.runs:
            src_run = src_para.runs[0]
            if src_run.font.size:
                new_run.font.size = src_run.font.size
            if src_run.font.name:
                new_run.font.name = src_run.font.name
    except Exception:
        pass



# -------------------------------------------------------
# Main: fill FLB template
# -------------------------------------------------------
async def fill_flb_document(template_path: str, user_inputs: dict, cover_keys: list = ["Projekt", "Objektadresse", "Ansprechpartner"]):
    doc = Document(template_path)

    # -------------------------------------------------------
    # STEP 1 — Fill cover sheet table fields (copy formatting)
    # -------------------------------------------------------

    # Try to find the table which contains the labels
    # We'll scan all tables and try to match keys in left column
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            # normalize key (remove colon and whitespace)
            left_text = row.cells[0].text.strip().replace(":", "").strip()

            if left_text in cover_keys and left_text in user_inputs:
                set_cell_text_preserve_style(row.cells[0], row.cells[1], user_inputs[left_text])

    # -------------------------------------------------------
    # STEP 2 — Detect sections by Heading 2
    # -------------------------------------------------------
    paragraphs = doc.paragraphs
    sections = []
    current_section = None

    for para in paragraphs:
        # Guard: some documents might have style names localized; ensure exact match to your template
        if para.style.name == "Heading 2":
            if current_section:
                sections.append(current_section)
            current_section = {
                "title_para": para,
                "instruction_paras": []
            }
        else:
            if current_section:
                current_section["instruction_paras"].append(para)

    if current_section:
        sections.append(current_section)

    # -------------------------------------------------------
    # STEP 3 — Generate and insert section texts
    # -------------------------------------------------------
    rag = OpenAIRAG(
            projektbezeichnung=user_inputs["Projekt"],
            objektadresse=user_inputs["Objektadresse"],
            ansprechpartner=user_inputs["Ansprechpartner"],
        )
    generated_sections = []
    for section in sections:
        instructions = "\n".join(p.text for p in section["instruction_paras"]).strip()
        instructions_clean = normalize_instructions(instructions)
        print(instructions_clean)
        contextual_prompt = rag.build_prompt_with_context(instructions_clean, generated_sections, user_inputs["Projekt"])
        generated_section = await rag.generate_section(contextual_prompt)
        generated_sections.append(generated_section)
        # Completely remove all instruction paragraphs (including bullets)
        for p in section["instruction_paras"]:
            delete_paragraph(p)

        insert_text_block_after(section["title_para"], generated_section.content)
    # -------------------------------------------------------
    # STEP 4 — Save output (template untouched)
    # -------------------------------------------------------
    #doc.save(output_path)
    #print(f"Created document: {output_path}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()