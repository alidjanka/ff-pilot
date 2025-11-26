import streamlit as st
from docx import Document
from io import BytesIO

from RAG.openai_rag import OpenAIRAG

import datetime
import asyncio

from prompts.prompts import Sections

async def generate_document(sections):
    rag = OpenAIRAG(collection_name="ff-pilot")
    full_doc = await rag.build_document_text(sections)

    metadata = f"""---
    title: "FLB Dokument"
    date: "{datetime.date.today()}"
    ---

    """

    full_doc_with_metadata = metadata + full_doc
    #with open('output/doc.md', "w", encoding="utf-8") as f:
    #    f.write(full_doc_with_metadata)
    return full_doc

def ask(query):
    rag = OpenAIRAG(collection_name="ff-pilot")
    answer, file_names = rag.query(query)
    return answer, file_names

def md_to_docx(md_text: str) -> bytes:
    """
    Convert markdown or plain text into a .docx file (simple formatting).
    Returns raw bytes, ready for Streamlit download.
    """
    doc = Document()
    lines = md_text.split("\n")

    for line in lines:
        if line.startswith("# "):        # H1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):     # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):    # H3
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

st.set_page_config(page_title="Dokument Generator", layout="wide")

# ------------------- Initialize session state -------------------
if "generated_doc" not in st.session_state:
    st.session_state.generated_doc = None

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ------------------- Layout -------------------
st.title("Mogli - FF Pilot")

# Two columns for split screen
left_col, right_col = st.columns([2, 1])  # bigger left side


# ------------------- LEFT COLUMN: DOCUMENT GENERATOR -------------------
with left_col:
    st.header("📄 FLB Generator")

    if st.button("📄 Dokument mit Vorlage generieren"):     
        st.session_state.generated_doc = asyncio.run(generate_document(Sections))
        st.success("Dokument wurde generiert!")

    st.markdown("---")

    if st.session_state.generated_doc:
        st.subheader("📘 Generiertes Dokument")
        st.markdown(st.session_state.generated_doc)
        docx_bytes = md_to_docx(st.session_state.generated_doc)
        st.download_button(
        label="📥 Download",
        data=docx_bytes,
        file_name=f"FLB_Repowering_{datetime.date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Noch kein Dokument generiert.")

# ------------------- RIGHT COLUMN: CHAT INTERFACE -------------------
with right_col:
    st.header("💬 Chat mit deinen Unterlagen")

    # Show chat history
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    # Chat input at bottom
    user_input = st.chat_input("Nachricht eingeben...")

    if user_input:
        # Append user message
        st.session_state.chat_history.append({
            "role": "user",
            "content": user_input
        })

        answer, file_names = ask(user_input)

        st.session_state.chat_history.append({
            "role": "assistant",
            "content": f'''
            Gemäß den folgenden Dateien: {file_names}: \n\n 
            {answer}
            '''
        })

        st.rerun()